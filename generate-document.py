import os
from docx import Document
from docx.shared import Pt

def create_code_doc(output_filename="CortexAI_SourceCode.docx"):
    doc = Document()
    doc.add_heading('CortexAI Source Code', 0)

    # Which files to include and folders to ignore
    valid_extensions = {'.py', '.yaml', '.ts', '.tsx', '.json', '.md'}
    ignore_dirs = {'venv', '.git', '__pycache__', 'node_modules', 'data', 'artifacts'}

    for root, dirs, files in os.walk('.'):
        # Skip ignored directories
        dirs[:] = [d for d in dirs if d not in ignore_dirs]
        
        for file in files:
            if os.path.splitext(file)[1] in valid_extensions:
                file_path = os.path.join(root, file)
                
                # Add file path as a heading
                doc.add_heading(file_path, level=2)
                
                # Read and add code
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        code = f.read()
                    
                    p = doc.add_paragraph(code)
                    # Set font to a monospace code font
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(8)
                except Exception as e:
                    doc.add_paragraph(f"Could not read {file_path}: {e}")

    doc.save(output_filename)
    print(f"✅ Successfully created {output_filename}")

if __name__ == "__main__":
    create_code_doc()